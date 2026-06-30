import io
from flask import request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from database import db
from models import Resume
from . import api_bp

@api_bp.route('/resumes', methods=['GET'])
@jwt_required()
def get_resumes():
    user_id = int(get_jwt_identity())
    resumes = Resume.query.filter_by(user_id=user_id).order_by(Resume.updated_at.desc()).all()
    return jsonify([r.to_dict() for r in resumes]), 200

@api_bp.route('/resumes/<int:resume_id>', methods=['GET'])
@jwt_required()
def get_resume(resume_id):
    user_id = int(get_jwt_identity())
    resume = Resume.query.filter_by(id=resume_id, user_id=user_id).first()
    if not resume:
        return jsonify({'message': 'Resume not found'}), 404
    return jsonify(resume.to_dict()), 200

@api_bp.route('/resumes', methods=['POST'])
@jwt_required()
def create_resume():
    user_id = int(get_jwt_identity())
    data = request.get_json() or {}
    title = data.get('title', 'Untitled Resume')
    template_name = data.get('template_name', 'modern')
    
    # Default blank structure matching 9 steps
    default_content = {
        'personal': {
            'name': '', 'email': '', 'phone': '', 'address': '',
            'linkedin': '', 'github': '', 'portfolio': '', 'summary': ''
        },
        'education': [],
        'experience': [],
        'skills': [],
        'projects': [],
        'certificates': [],
        'languages': [],
        'achievements': []
    }
    
    content = data.get('content', default_content)
    custom_styling = data.get('custom_styling', {
        'font_family': 'Inter',
        'font_size': 'medium',
        'spacing': 'normal',
        'theme_color': '#1e3a8a'
    })
    
    resume = Resume(
        user_id=user_id,
        title=title,
        template_name=template_name,
        content=content,
        custom_styling=custom_styling,
        ats_score=0
    )
    
    db.session.add(resume)
    db.session.commit()
    return jsonify(resume.to_dict()), 201

@api_bp.route('/resumes/<int:resume_id>', methods=['PUT'])
@jwt_required()
def update_resume(resume_id):
    user_id = int(get_jwt_identity())
    resume = Resume.query.filter_by(id=resume_id, user_id=user_id).first()
    if not resume:
        return jsonify({'message': 'Resume not found'}), 404
        
    data = request.get_json() or {}
    if 'title' in data:
        resume.title = data['title']
    if 'template_name' in data:
        resume.template_name = data['template_name']
    if 'content' in data:
        resume.content = data['content']
    if 'custom_styling' in data:
        resume.custom_styling = data['custom_styling']
    if 'ats_score' in data:
        resume.ats_score = data['ats_score']
        
    db.session.commit()
    return jsonify(resume.to_dict()), 200

@api_bp.route('/resumes/<int:resume_id>/clone', methods=['POST'])
@jwt_required()
def clone_resume(resume_id):
    user_id = int(get_jwt_identity())
    resume = Resume.query.filter_by(id=resume_id, user_id=user_id).first()
    if not resume:
        return jsonify({'message': 'Resume not found'}), 404
        
    cloned = Resume(
        user_id=user_id,
        title=f"Copy of {resume.title}",
        template_name=resume.template_name,
        content=resume.content,
        custom_styling=resume.custom_styling,
        ats_score=resume.ats_score
    )
    
    db.session.add(cloned)
    db.session.commit()
    return jsonify(cloned.to_dict()), 201

@api_bp.route('/resumes/<int:resume_id>', methods=['DELETE'])
@jwt_required()
def delete_resume(resume_id):
    user_id = int(get_jwt_identity())
    resume = Resume.query.filter_by(id=resume_id, user_id=user_id).first()
    if not resume:
        return jsonify({'message': 'Resume not found'}), 404
        
    db.session.delete(resume)
    db.session.commit()
    return jsonify({'message': 'Resume deleted successfully'}), 200

# --- Server Side DOCX Exporter ---

@api_bp.route('/resumes/<int:resume_id>/export-docx', methods=['GET'])
@jwt_required()
def export_docx(resume_id):
    user_id = int(get_jwt_identity())
    resume = Resume.query.filter_by(id=resume_id, user_id=user_id).first()
    if not resume:
        return jsonify({'message': 'Resume not found'}), 404
        
    content = resume.content
    personal = content.get('personal', {})
    education = content.get('education', [])
    experience = content.get('experience', [])
    skills = content.get('skills', [])
    projects = content.get('projects', [])
    certificates = content.get('certificates', [])
    languages = content.get('languages', [])
    achievements = content.get('achievements', [])
    
    # Create python-docx document
    doc = Document()
    
    # Adjust page margins (0.75" is typical for professional resumes)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Styles config
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Theme color (Default: deep navy)
    primary_color = RGBColor(0x1E, 0x3A, 0x8A)
    
    # 1. Header (Name, Contacts)
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(personal.get('name', 'Your Name').upper())
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = primary_color
    
    # Contacts details line
    contacts = []
    if personal.get('email'): contacts.append(personal['email'])
    if personal.get('phone'): contacts.append(personal['phone'])
    if personal.get('address'): contacts.append(personal['address'])
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(12)
    contact_run = contact_p.add_run(" | ".join(contacts))
    contact_run.font.size = Pt(9.5)
    
    # Social links
    socials = []
    if personal.get('linkedin'): socials.append(f"LinkedIn: {personal['linkedin']}")
    if personal.get('github'): socials.append(f"GitHub: {personal['github']}")
    if personal.get('portfolio'): socials.append(f"Portfolio: {personal['portfolio']}")
    
    if socials:
        social_p = doc.add_paragraph()
        social_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        social_p.paragraph_format.space_after = Pt(16)
        social_run = social_p.add_run(" | ".join(socials))
        social_run.font.size = Pt(9.5)
        social_run.font.italic = True
        
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = primary_color
        
        # Add a subtle bottom border equivalent (using text underlining or line separation)
        # For simplicity in python-docx, we can add a horizontal rule or border paragraph
        # doc.add_paragraph("―" * 50)
        
    # 2. Professional Summary
    if personal.get('summary'):
        add_section_heading("Professional Summary")
        summary_p = doc.add_paragraph(personal['summary'])
        summary_p.paragraph_format.space_after = Pt(12)
        summary_p.paragraph_format.line_spacing = 1.15

    # 3. Work Experience
    if experience:
        add_section_heading("Professional Experience")
        for exp in experience:
            exp_p = doc.add_paragraph()
            exp_p.paragraph_format.space_after = Pt(2)
            role_run = exp_p.add_run(exp.get('role', 'Job Title'))
            role_run.bold = True
            
            comp_text = f", {exp.get('company', 'Company Name')}"
            exp_p.add_run(comp_text)
            
            # Date alignment
            date_p = doc.add_paragraph()
            date_p.paragraph_format.space_after = Pt(4)
            date_text = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
            if exp.get('location'):
                date_text += f" | {exp['location']}"
            date_run = date_p.add_run(date_text)
            date_run.font.italic = True
            date_run.font.size = Pt(9.5)
            
            # Bullet points
            desc = exp.get('description', '')
            if desc:
                # Support single string with newlines or arrays
                bullets = desc.split('\n') if '\n' in desc else [desc]
                for bullet in bullets:
                    if bullet.strip():
                        # Clean leading bullet symbol if it exists
                        b_text = re.sub(r'^[•\-\*]\s*', '', bullet.strip())
                        doc.add_paragraph(b_text, style='List Bullet')
                        
            # Space between roles
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            
    # 4. Education
    if education:
        add_section_heading("Education")
        for edu in education:
            edu_p = doc.add_paragraph()
            edu_p.paragraph_format.space_after = Pt(2)
            deg_run = edu_p.add_run(edu.get('degree', 'Degree'))
            deg_run.bold = True
            
            school_text = f", {edu.get('school', 'School Name')}"
            edu_p.add_run(school_text)
            
            date_p = doc.add_paragraph()
            date_p.paragraph_format.space_after = Pt(4)
            date_text = f"{edu.get('start_date', '')} – {edu.get('end_date', '')}"
            if edu.get('location'):
                date_text += f" | {edu['location']}"
            date_run = date_p.add_run(date_text)
            date_run.font.italic = True
            date_run.font.size = Pt(9.5)
            
            # Education details / GPA
            details = edu.get('description', '')
            if details:
                edu_details_p = doc.add_paragraph(details)
                edu_details_p.paragraph_format.space_after = Pt(6)

    # 5. Projects
    if projects:
        add_section_heading("Projects")
        for proj in projects:
            proj_p = doc.add_paragraph()
            proj_p.paragraph_format.space_after = Pt(2)
            title_run = proj_p.add_run(proj.get('title', 'Project Title'))
            title_run.bold = True
            
            if proj.get('link'):
                proj_p.add_run(f" ({proj['link']})")
                
            desc_p = doc.add_paragraph(proj.get('description', ''))
            desc_p.paragraph_format.space_after = Pt(6)

    # 6. Skills
    if skills:
        add_section_heading("Skills")
        skills_p = doc.add_paragraph()
        skills_p.paragraph_format.space_after = Pt(8)
        skills_p.add_run(", ".join(skills) if isinstance(skills, list) else skills)

    # 7. Certifications & Extras
    if certificates:
        add_section_heading("Certifications")
        for cert in certificates:
            cert_p = doc.add_paragraph()
            cert_p.paragraph_format.space_after = Pt(2)
            title_run = cert_p.add_run(cert.get('title', 'Certificate Title'))
            title_run.bold = True
            
            issuer_text = ""
            if cert.get('issuer'):
                issuer_text += f" - {cert['issuer']}"
            if cert.get('date'):
                issuer_text += f" ({cert['date']})"
            cert_p.add_run(issuer_text)

    # 8. Languages
    if languages:
        add_section_heading("Languages")
        langs = []
        for lang in languages:
            langs.append(f"{lang.get('language', '')} ({lang.get('proficiency', 'Fluent')})")
        lang_p = doc.add_paragraph(", ".join(langs))
        lang_p.paragraph_format.space_after = Pt(8)

    # 9. Achievements
    if achievements:
        add_section_heading("Key Achievements")
        for ach in achievements:
            doc.add_paragraph(ach.get('description', '') or ach.get('title', ''), style='List Bullet')
            
    # Save to a memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    filename = f"{personal.get('name', 'resume').replace(' ', '_')}_resume.docx"
    
    return send_file(
        file_stream,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename
    )

# --- Public Sharing Endpoints ---

@api_bp.route('/resumes/public/<int:resume_id>', methods=['GET'])
def get_public_resume(resume_id):
    resume = Resume.query.get(resume_id)
    if not resume:
        return jsonify({'message': 'Resume not found'}), 404
        
    if not resume.is_public:
        return jsonify({'message': 'This resume is private. Owner must enable sharing.'}), 403
        
    return jsonify(resume.to_dict()), 200

@api_bp.route('/resumes/<int:resume_id>/toggle-public', methods=['PUT'])
@jwt_required()
def toggle_resume_public(resume_id):
    user_id = int(get_jwt_identity())
    resume = Resume.query.filter_by(id=resume_id, user_id=user_id).first()
    if not resume:
        return jsonify({'message': 'Resume not found'}), 404
        
    resume.is_public = not resume.is_public
    db.session.commit()
    
    return jsonify({
        'message': f"Resume is now {'public' if resume.is_public else 'private'}",
        'is_public': resume.is_public
    }), 200

